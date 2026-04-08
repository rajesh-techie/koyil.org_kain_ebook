"""
scraper.py — URL to DOCX Web Scraper
Requires: playwright, beautifulsoup4, lxml, python-docx
Usage:
    python scraper.py --url "https://example.com" --output "C:\\output\\result.docx"
"""

import argparse
import json
import logging
import os
import re
import subprocess
import sys
from pathlib import Path
from urllib.parse import urlparse, urljoin, urldefrag
from urllib.request import urlopen
from io import BytesIO

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from playwright.sync_api import sync_playwright, TimeoutError as PlaywrightTimeout

# ---------------------------------------------------------------------------
# Logging setup
# ---------------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.StreamHandler(sys.stdout)],
)
log = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Config loader
# ---------------------------------------------------------------------------
CONFIG_PATH = Path(__file__).parent.parent / "config.json"


def load_config() -> dict:
    """Load configuration from config.json next to this script."""
    if not CONFIG_PATH.exists():
        log.error("config.json not found at: %s", CONFIG_PATH)
        sys.exit(1)
    with CONFIG_PATH.open("r", encoding="utf-8") as f:
        cfg = json.load(f)
    log.info("Config loaded from %s", CONFIG_PATH)
    return cfg


# ---------------------------------------------------------------------------
# Input validation (Security: T7.1, T7.2)
# ---------------------------------------------------------------------------
_ALLOWED_SCHEMES = {"http", "https"}


def validate_url(url: str) -> str:
    """Validate URL — must be http/https with a non-empty host."""
    try:
        parsed = urlparse(url)
    except Exception:
        log.error("URL parsing failed: %s", url)
        sys.exit(1)
    if parsed.scheme not in _ALLOWED_SCHEMES:
        log.error("URL scheme '%s' not allowed. Only http/https permitted.", parsed.scheme)
        sys.exit(1)
    if not parsed.netloc:
        log.error("URL has no host: %s", url)
        sys.exit(1)
    # Prevent localhost/private IP access (basic SSRF guard)
    host = parsed.hostname or ""
    if host in ("localhost", "127.0.0.1", "0.0.0.0", "::1"):
        log.error("URL points to a local address — not permitted.")
        sys.exit(1)
    return url


def validate_output_path(raw_path: str) -> Path:
    """Validate output file path — must end in .docx, no traversal."""
    path = Path(raw_path).resolve()
    # Prevent directory traversal by ensuring resolve() stays under a real location
    if ".." in Path(raw_path).parts:
        log.error("Output path contains '..': %s", raw_path)
        sys.exit(1)
    if path.suffix.lower() != ".docx":
        log.error("Output file must have a .docx extension. Got: %s", path.suffix)
        sys.exit(1)
    # Ensure the parent directory exists (or can be created)
    parent = path.parent
    if not parent.exists():
        try:
            parent.mkdir(parents=True, exist_ok=True)
            log.info("Created output directory: %s", parent)
        except OSError as exc:
            log.error("Cannot create output directory '%s': %s", parent, exc)
            sys.exit(1)
    # Quick write-permission check
    if not os.access(parent, os.W_OK):
        log.error("Output directory is not writable: %s", parent)
        sys.exit(1)
    return path


# ---------------------------------------------------------------------------
# CLI argument parsing
# ---------------------------------------------------------------------------
def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Scrape a web page and save structured content as a .docx file."
    )
    parser.add_argument(
        "--url",
        required=True,
        help="Full URL of the page to scrape (http/https only).",
    )
    parser.add_argument(
        "--output",
        required=True,
        help="Full file path for the output document (must end in .docx).",
    )
    parser.add_argument(
        "--title",
        required=True,
        help="Document title to be used in header and TOC.",
    )
    parser.add_argument(
        "--language",
        default="english",
        help="Language for template selection (english, tamil, hindi, telugu, malayalam, kannada). Default: english",
    )
    return parser.parse_args()


# ---------------------------------------------------------------------------
# Browser launch & navigation (Task 2)
# ---------------------------------------------------------------------------
def open_page(playwright, url: str, cfg: dict):
    """
    Launch a Chromium browser, open a new page, navigate to url.
    Returns (browser, page) — caller is responsible for closing browser.
    """
    # Use reading_mode setting if available, otherwise fall back to headless setting
    reading_mode = cfg.get("reading_mode", True)
    headless = cfg.get("browser", {}).get("headless", False)
    
    # In reading mode, enable headless for cleaner content extraction
    if reading_mode:
        headless = True
    
    timeout_ms = cfg.get("browser", {}).get("timeout_ms", 30000)

    log.info("Launching Chromium (headless=%s) ...", headless)
    browser = playwright.chromium.launch(headless=headless)
    context = browser.new_context()
    page = context.new_page()
    page.set_default_timeout(timeout_ms)

    log.info("Navigating to: %s", url)
    try:
        page.goto(url, wait_until="networkidle")
        log.info("Page loaded: %s", page.title())
    except PlaywrightTimeout:
        log.error("Timed out loading page: %s", url)
        browser.close()
        sys.exit(1)

    return browser, page


# ---------------------------------------------------------------------------
# Content extraction helpers (Task 3)
# ---------------------------------------------------------------------------

# Semantic block-level tags — each one becomes its own paragraph.
# div/span deliberately excluded: they are containers that cause content duplication.
_BLOCK_CONTENT_TAGS = {"p", "li", "h1", "h2", "h3", "h4", "h5", "h6",
                        "td", "th", "blockquote", "pre"}

# Tags that map to a Word heading style (level capped at 4 so they don't
# clash with the document's own Heading 1/2 used for chapters).
_HEADING_STYLE = {"h1": "Heading 3", "h2": "Heading 3",
                  "h3": "Heading 3", "h4": "Heading 4",
                  "h5": "Heading 4", "h6": "Heading 4"}

# Tags that are purely structural/navigation noise — stripped entirely
_STRIP_TAGS = {"script", "style", "noscript", "nav", "footer", "header",
               "aside", "form", "button", "iframe"}


def get_page_html(page) -> str:
    """Return the full rendered HTML from the Playwright page object."""
    return page.content()


def parse_lines(html: str, page_url: str = "", include_images: bool = False) -> list[dict]:
    """
    Parse rendered HTML into structured items with preserved formatting and inline images.
    Returns list of:
    - Text items: {"style": str, "runs": [...], "type": "text"}
    - Image items: {"type": "image", "url": str, "alt": str}
    
    Images are preserved at their original inline positions in the content.
    """
    soup = BeautifulSoup(html, "lxml")

    for tag in soup.find_all(_STRIP_TAGS):
        tag.decompose()

    body = (
        soup.find("article")
        or soup.find("main")
        or soup.find("div", {"class": re.compile(r"(content|entry|post|article|body)", re.I)})
        or soup.find("body")
    )
    if body is None:
        body = soup

    def extract_runs(element):
        """
        Walk element tree and extract list of {"text": str, "bold": bool, "italic": bool}.
        Applies formatting based on parent tags.
        Preserves word spacing by not stripping whitespace from text nodes.
        Skips problematic formatting tags like underline, strikethrough, mark, etc.
        """
        # Tags to skip (these don't map to standard markdown and can cause formatting issues)
        _SKIP_TAGS = {"u", "s", "strike", "del", "mark", "sup", "sub", "abbr", "code", "kbd"}
        
        def walk_text(node, bold=False, italic=False, is_root=False):
            """Recursively walk node, collecting text with formatting."""
            runs = []
            if isinstance(node, str):
                # Preserve original spacing
                text = str(node)
                # Only skip if entirely whitespace
                if text.strip():
                    # Collapse multiple whitespace to single space
                    text = " ".join(text.split())
                    if text:  # double-check after collapse
                        runs.append({"text": text, "bold": bold, "italic": italic})
                return runs
            
            # If node is a tag, check if it's inline formatting
            if hasattr(node, 'name'):
                # Don't descend into nested block elements (unless it's the root)
                if not is_root and node.name in _BLOCK_CONTENT_TAGS:
                    return runs
                
                # Skip problematic formatting tags but extract their text
                if node.name in _SKIP_TAGS:
                    # Extract text from skip tags without applying their formatting
                    for child in node.children:
                        runs.extend(walk_text(child, bold, italic, is_root=False))
                    return runs
                
                # Apply formatting for standard inline tags
                new_bold = bold or (node.name in {"b", "strong"})
                new_italic = italic or (node.name in {"i", "em"})
                
                # Walk children
                for child in node.children:
                    runs.extend(walk_text(child, new_bold, new_italic, is_root=False))
            
            return runs
        
        return walk_text(element, is_root=True)

    items = []
    seen = set()

    def process_node(node):
        """Process a node, adding text items and image items in document order."""
        nonlocal items, seen
        
        if isinstance(node, str):
            return
        
        if not hasattr(node, 'name'):
            return
        
        # If this is a block content element, add it as a text item
        if node.name in _BLOCK_CONTENT_TAGS:
            # Skip container elements with block children
            if any(getattr(child, "name", None) in _BLOCK_CONTENT_TAGS
                   for child in node.children):
                # Still process children of containers
                for child in node.children:
                    process_node(child)
                return
            
            # Check if this block element contains <br> tags (which should split into separate items)
            br_tags = list(node.find_all("br", recursive=False))  # Only direct BR children
            if br_tags:
                # Split content by direct <br> tags and create separate items
                current_part = []
                
                for child in list(node.children):  # Use list() to avoid iterator issues
                    if hasattr(child, 'name') and child.name == "br":
                        # <br> tag encountered, finalize current part
                        if current_part:
                            # Build runs directly from the parts without moving nodes
                            runs = []
                            for part in current_part:
                                if isinstance(part, str):
                                    text = str(part).strip()
                                    if text:
                                        text = " ".join(text.split())
                                        if text:
                                            runs.append({"text": text, "bold": False, "italic": False})
                                else:
                                    # Extract runs from tag without moving it
                                    part_runs = extract_runs(part)
                                    runs.extend(part_runs)
                            
                            if runs:
                                full_text = " ".join(r["text"] for r in runs)
                                if full_text not in seen:
                                    seen.add(full_text)
                                    
                                    if node.name in _HEADING_STYLE:
                                        style = _HEADING_STYLE[node.name]
                                    elif node.name == "li":
                                        parent = node.parent
                                        if parent and parent.name == "ol":
                                            style = "List Number"
                                        else:
                                            style = "List Bullet"
                                    else:
                                        style = "Normal"
                                    
                                    items.append({"style": style, "runs": runs, "type": "text"})
                            current_part = []
                    else:
                        current_part.append(child)
                
                # Process remaining part after last <br>
                if current_part:
                    runs = []
                    for part in current_part:
                        if isinstance(part, str):
                            text = str(part).strip()
                            if text:
                                text = " ".join(text.split())
                                if text:
                                    runs.append({"text": text, "bold": False, "italic": False})
                        else:
                            part_runs = extract_runs(part)
                            runs.extend(part_runs)
                    
                    if runs:
                        full_text = " ".join(r["text"] for r in runs)
                        if full_text not in seen:
                            seen.add(full_text)
                            
                            if node.name in _HEADING_STYLE:
                                style = _HEADING_STYLE[node.name]
                            elif node.name == "li":
                                parent = node.parent
                                if parent and parent.name == "ol":
                                    style = "List Number"
                                else:
                                    style = "List Bullet"
                            else:
                                style = "Normal"
                            
                            items.append({"style": style, "runs": runs, "type": "text"})
                return
            
            # Normal processing without <br> tags
            runs = extract_runs(node)
            if runs:
                full_text = " ".join(r["text"] for r in runs)
                if full_text not in seen:
                    seen.add(full_text)
                    
                    if node.name in _HEADING_STYLE:
                        style = _HEADING_STYLE[node.name]
                    elif node.name == "li":
                        parent = node.parent
                        if parent and parent.name == "ol":
                            style = "List Number"
                        else:
                            style = "List Bullet"
                    else:
                        style = "Normal"
                    
                    items.append({"style": style, "runs": runs, "type": "text"})
        
        # If this is an image tag and images are enabled, add it as an image item
        elif node.name == "img" and include_images:
            src = node.get("src", "").strip()
            if src and not src.startswith("data:"):
                abs_url = urljoin(page_url, src)
                alt = node.get("alt", "Image").strip() or "Image"
                items.append({
                    "type": "image",
                    "url": abs_url,
                    "alt": alt
                })
                log.debug("Found image: %s (%s)", alt, abs_url)
        
        # Recursively process children to maintain document order
        for child in node.children:
            process_node(child)

    # Process all children of body to maintain document order
    for child in body.children:
        process_node(child)

    return items


def _line_matches_any(line: str, markers: list[str]) -> bool:
    """Return True if any marker string appears (case-insensitive) in line."""
    line_lower = line.lower()
    return any(m.lower() in line_lower for m in markers)


def extract_content(items: list[dict], start_markers: list[str], stop_markers: list[str], remove_text: list[str] = None, parameter1: str = "") -> list[dict]:
    """
    Return the slice of items between the start and stop markers.

    Start rule : find the first item whose text contains any start_marker.
                 Include the item BEFORE that match (one item prior).
    Stop rule  : find the first item (after start) whose text contains any
                 stop_marker. Include the stop marker line itself.

    Then filter out any items whose full text matches any remove_text entries (case-insensitive).
    Then append a new "Content Source" line with parameter1 (URL).

    Returns an empty list if start marker is not found.
    """
    if remove_text is None:
        remove_text = []

    # Build full-text map for marker matching
    texts = [" ".join(r.get("text", "") for r in item.get("runs", []))
             for item in items]

    start_idx = None
    for i, text in enumerate(texts):
        if _line_matches_any(text, start_markers):
            start_idx = max(0, i - 1)
            log.info("Start marker found at item %d: %r", i, text[:80])
            break

    if start_idx is None:
        log.warning("No start marker found. Markers tried: %s", start_markers)
        return []

    stop_idx = None
    for i, text in enumerate(texts[start_idx:], start=start_idx):
        if i == start_idx:
            continue  # never match stop on the very first item
        if _line_matches_any(text, stop_markers):
            stop_idx = i + 1  # Slice is exclusive, so this includes the marker itself
            log.info("Stop marker found at item %d: %r  — including it",
                     i, text[:80])
            break

    if stop_idx is None:
        log.warning("No stop marker found — extracting to end of page. Markers tried: %s",
                    stop_markers)
        stop_idx = len(items)

    extracted = items[start_idx:stop_idx]
    
    # Filter out remove_text items
    if remove_text:
        remove_text_lower = [t.lower() for t in remove_text]
        filtered = []
        for item in extracted:
            full_text = " ".join(r.get("text", "") for r in item.get("runs", []))
            # Skip if this item's text matches any remove_text (exact match, case-insensitive)
            if not any(full_text.lower() == rt for rt in remove_text_lower):
                filtered.append(item)
        extracted = filtered
        log.info("After filtering remove_text: %d items remaining (was %d).", len(extracted), len(extracted) + (len([i for i in items[start_idx:stop_idx] if " ".join(r.get("text", "") for r in i.get("runs", [])).lower() in remove_text_lower])))
    
    # Append "Content Source" line with parameter1 (URL) if provided
    if parameter1:
        extracted.append({
            "style": "Normal",
            "runs": [{"text": f"Content Source - {parameter1}", "bold": False, "italic": False}],
            "type": "text"
        })
        log.info("Added Content Source line: %s", parameter1)
    
    log.info("Extracted %d items from page.", len(extracted))
    return extracted


def scrape_page(page, cfg: dict, page_url: str = "") -> list[dict]:
    """
    Full pipeline for a single page:
      1. Get rendered HTML
      2. Parse into structured items (Reader Mode simulation) with inline images
      3. Apply start/stop marker extraction
      4. Remove static text items
    Returns list of items (text items with "runs" and image items with "url"/"alt").
    """
    start_markers = cfg.get("start_markers", [])
    stop_markers = cfg.get("stop_markers", [])
    remove_text = cfg.get("remove_text", [])
    include_images = cfg.get("include_images", False)

    if not start_markers:
        log.warning("No start_markers defined in config.json")
    if not stop_markers:
        log.warning("No stop_markers defined in config.json")

    html = get_page_html(page)
    items = parse_lines(html, page_url, include_images)
    log.info("Parsed %d total items from page HTML.", len(items))

    if not items:
        log.warning("Page produced no text — check URL or page structure.")
        return []

    return extract_content(items, start_markers, stop_markers, remove_text, page_url)


# ---------------------------------------------------------------------------
# Chapter / TOC detection (Task 4)
# ---------------------------------------------------------------------------

def _get_content_root(soup: BeautifulSoup):
    """Return the semantic content root element (same strategy as parse_lines)."""
    return (
        soup.find("article")
        or soup.find("main")
        or soup.find("div", {"class": re.compile(r"(content|entry|post|article|body)", re.I)})
        or soup.find("body")
        or soup
    )


def detect_chapters(html: str, base_url: str, toc_selector: str, max_chapters: int) -> list[dict]:
    """
    Detect chapter links from the page HTML robustly.

    Strategy:
      1. Scope to the semantic content area (article/main) — excludes nav/header/footer.
      2. Select all elements matching toc_selector (default: "li > a[href]").
      3. Resolve every href to an absolute URL using base_url.
      4. Keep only same-domain links (avoids external/social links).
      5. Deduplicate by normalised URL (fragment stripped).
      6. Cap at max_chapters.

    Returns a list of dicts: [{"title": str, "url": str}, ...]
    Title = link text, whatever it is — no hardcoding.
    """
    base_parsed = urlparse(base_url)
    base_domain = base_parsed.netloc.lower()

    # Parse a clean copy (don't mutate the one used for text extraction)
    soup = BeautifulSoup(html, "lxml")

    # Remove noise containers before scanning
    for tag in soup.find_all(_STRIP_TAGS):
        tag.decompose()

    content_root = _get_content_root(soup)

    chapters = []
    seen_urls = set()

    for a_tag in content_root.select(toc_selector):
        href = a_tag.get("href", "").strip()
        if not href or href.startswith(("#", "javascript:", "mailto:")):
            continue

        # Resolve relative URLs → absolute
        abs_url, _ = urldefrag(urljoin(base_url, href))

        # Same-domain filter (robust: works for any site)
        parsed = urlparse(abs_url)
        if parsed.netloc.lower() != base_domain:
            log.debug("Skipping external link: %s", abs_url)
            continue

        # Deduplicate
        norm = abs_url.rstrip("/").lower()
        if norm in seen_urls:
            continue
        seen_urls.add(norm)

        # Title = whatever text is in the link — no assumptions
        title = a_tag.get_text(separator=" ", strip=True)
        if not title:
            title = abs_url  # fallback: use URL if link has no visible text
        
        # Clean up chapter title: remove common prefix characters (dashes, bullets, symbols)
        # These often appear as HTML list markers or separators
        title = title.lstrip("–—-•·‣»›:[] ").strip()
        if not title:
            title = abs_url  # fallback if title becomes empty after cleaning
        
        # Skip self-referencing links (Table of Contents, Introduction, current page URL)
        title_lower = title.lower()
        if title_lower in ("table of contents", "introduction"):
            log.debug("Skipping self-reference: %s", title)
            continue
        if abs_url.rstrip("/").lower() == base_url.rstrip("/").lower():
            log.debug("Skipping current page link: %s", title)
            continue

        chapters.append({"title": title, "url": abs_url})

        if len(chapters) >= max_chapters:
            log.warning("Reached max_chapters cap (%d). Additional links ignored.", max_chapters)
            break

    log.info("Detected %d chapter link(s) on page.", len(chapters))
    return chapters


# ---------------------------------------------------------------------------
# DOCX generation (Task 5)
# ---------------------------------------------------------------------------

def _add_toc_field(doc: Document):
    """
    Insert a Word TOC field code (\"Table of Contents\") so that when the .docx
    is opened in MS Word and the user presses Ctrl+A then F9 (Update Fields),
    it renders with real page numbers. This is the standard Word TOC approach.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr_run = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    instr_run._r.append(instr_text)

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    run2 = paragraph.add_run()
    run2._r.append(fld_char_sep)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run3 = paragraph.add_run()
    run3._r.append(fld_char_end)


def _add_page_break(doc: Document):
    """Insert a hard page break."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _clear_all_headers_footers(doc: Document) -> None:
    """Clear all headers and footers from all sections in the document."""
    for section in doc.sections:
        # Clear header
        header = section.header
        for para in list(header.paragraphs):
            p = para._element
            p.getparent().remove(p)
        
        # Clear footer
        footer = section.footer
        for para in list(footer.paragraphs):
            p = para._element
            p.getparent().remove(p)


def _add_header(doc: Document, title: str = ""):
    """
    Add header with document title (parameter2) in blue color.
    
    Args:
        doc: Document object
        title: Document title to display in header
    """
    section = doc.sections[0]
    header = section.header
    
    # Clear any existing header content first
    for para in list(header.paragraphs):
        p = para._element
        p.getparent().remove(p)
    
    header_para = header.add_paragraph()
    header_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    if not title:
        title = "http://koyil.org"
    
    run = header_para.add_run(title)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 51, 204)  # Blue color


def _add_page_numbers(doc: Document, parameter1: str = "", page_url: str = ""):
    """
    Add footer with parameter1 (URL as hyperlink), page numbering, and koyil.org (blue color).
    Format: [parameter1-hyperlink] [centered page X] [right-koyil.org]
    
    Args:
        doc: Document object
        parameter1: URL of the source (left-aligned, as hyperlink)
        page_url: URL of the page being scraped (for legacy support)
    """
    section = doc.sections[0]
    footer = section.footer
    
    # Remove ALL existing paragraphs from footer - we'll create fresh ones
    while len(footer.paragraphs) > 0:
        p = footer.paragraphs[0]._element
        p.getparent().remove(p)
    
    # Create new footer paragraph from scratch
    footer_para = footer.add_paragraph()

    # Build footer with 3 sections: left | center page # | right
    # Tab stop at 3.25" for center, 6.5" for right
    footer_para.paragraph_format.tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
    footer_para.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)

    # Left: parameter1 (URL) as hyperlink in blue
    if parameter1:
        run_left_text = footer_para.add_run(parameter1)
        run_left_text.font.size = Pt(10)
        run_left_text.font.color.rgb = RGBColor(0, 51, 204)  # Blue
        run_left_text.font.underline = True  # Underline for hyperlink appearance
    else:
        run_left = footer_para.add_run("http://koyil.org")
        run_left.font.size = Pt(10)
        run_left.font.color.rgb = RGBColor(0, 51, 204)

    # Tab to center
    footer_para.add_run("\t")

    # Center: page number with field code
    run_page = footer_para.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run_page._r.append(fld_char_begin)

    instr_run = footer_para.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    instr_run._r.append(instr_text)

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    run_sep = footer_para.add_run()
    run_sep._r.append(fld_char_sep)

    # Page number placeholder
    run_num = footer_para.add_run("1")
    run_num.font.size = Pt(10)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_end = footer_para.add_run()
    run_end._r.append(fld_char_end)

    # Tab to right
    footer_para.add_run("\t")

    # Right: koyil.org in blue
    run_right = footer_para.add_run("koyil.org")
    run_right.font.size = Pt(10)
    run_right.font.color.rgb = RGBColor(0, 51, 204)  # Blue


def _make_bookmark_name(index: int, title: str) -> str:
    """Produce a valid Word bookmark name (XML NCName, max 40 chars)."""
    safe = re.sub(r"[^a-zA-Z0-9]", "_", title)
    safe = re.sub(r"_+", "_", safe).strip("_")
    return f"ch_{index}_{safe[:30]}"


def _add_bookmark(paragraph, bookmark_id: int, bookmark_name: str):
    """
    Wrap a paragraph in a Word bookmark so TOC hyperlinks can target it.
    bookmarkStart is inserted before the first run; bookmarkEnd at the end.
    """
    p = paragraph._p
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(bookmark_id))
    bm_start.set(qn("w:name"), bookmark_name)
    # Insert before the first child (before any run/pPr)
    p.insert(0, bm_start)

    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bookmark_id))
    p.append(bm_end)


def _add_toc_hyperlink(paragraph, text: str, bookmark_name: str):
    """
    Add an internal hyperlink (w:anchor) to an existing paragraph,
    styled with Word's built-in 'Hyperlink' character style.
    Clicking this in MS Word / Word Online navigates to the bookmark.
    """
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rpr.append(rStyle)
    run.append(rpr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _update_word_fields_async(docx_path: str):
    """
    Update all fields in Word document (PAGEREF fields for TOC page numbers).
    Uses PowerShell + COM on Windows. Properly releases file handle after closing.
    """
    if sys.platform != "win32":
        log.debug("Field update only supported on Windows. Skipping.")
        return
    
    docx_abs = str(Path(docx_path).resolve())
    
    # PowerShell script to update Word fields with proper handle cleanup
    ps_script = f"""
$doc_path = '{docx_abs}'
$word = $null
$doc = $null

try {{
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $word.ScreenUpdating = $false
    
    $doc = $word.Documents.Open($doc_path, $false, $false, $false)
    
    # Update all fields
    $fields = $doc.Fields
    foreach ($field in $fields) {{
        try {{ $field.Update() }} catch {{ }}
    }}
    
    # Save document
    $doc.Save()
}} catch {{
    Write-Host "ERROR:$_"
}} finally {{
    # Explicitly close and release all handles
    if ($doc -ne $null) {{
        $doc.Close($false)
        [System.Runtime.InteropServices.Marshal]::ReleaseComObject($doc) | Out-Null
    }}
    
    if ($word -ne $null) {{
        $word.Quit()
        [System.Runtime.InteropServices.Marshal]::ReleaseComObject($word) | Out-Null
    }}
    
    # Force garbage collection to release handles
    [GC]::Collect()
    [GC]::WaitForPendingFinalizers()
    
    # Small delay to ensure file handle is released
    Start-Sleep -Milliseconds 500
    
    Write-Host "SUCCESS"
}}
"""
    
    try:
        # Run PowerShell and wait for result
        result = subprocess.run(
            ["powershell", "-NoProfile", "-Command", ps_script],
            capture_output=True,
            timeout=300,
            text=True,
        )
        if "SUCCESS" in result.stdout:
            log.info("✓ Word fields updated successfully (TOC page numbers calculated)")
        else:
            log.debug("Word field update result: %s", result.stdout)
            log.info("  Tip: Open document and press Ctrl+A, then F9 to update TOC page numbers")
    except subprocess.TimeoutExpired:
        log.warning("Field update timed out (300s). You can manually update in Word: Ctrl+A, F9")
    except Exception as e:
        log.debug("Could not update Word fields: %s", e)
        log.info("  Tip: Open document and press Ctrl+A, then F9 to update TOC page numbers")


def _get_available_style(doc: Document, requested_style: str) -> str:
    """
    Get an available style from the document, with fallback mappings for common list styles.
    If the requested style doesn't exist, tries to find a suitable alternative.
    """
    try:
        # Check if style exists by trying to access it
        _ = doc.styles[requested_style]
        return requested_style
    except KeyError:
        pass
    
    # Map unavailable list styles to standard alternatives
    style_mappings = {
        "List Bullet": "List Bullet",  # Try the base version first
        "List Number": "List Number",
        "List Bullet 2": "List Bullet",
        "List Bullet 3": "List Bullet",
    }
    
    # Check if requested style is a list style that we should map
    if requested_style in style_mappings:
        mapped_style = style_mappings[requested_style]
        try:
            _ = doc.styles[mapped_style]
            log.debug("Mapped style '%s' to '%s'", requested_style, mapped_style)
            return mapped_style
        except KeyError:
            pass
    
    # Try standard list styles
    for fallback in ["List Bullet", "List Paragraph", "Normal"]:
        try:
            _ = doc.styles[fallback]
            log.debug("Falling back from '%s' to '%s'", requested_style, fallback)
            return fallback
        except KeyError:
            continue
    
    # Last resort: use Normal
    log.debug("No suitable style found for '%s'; using 'Normal'", requested_style)
    return "Normal"


def _write_lines(doc: Document, items: list[dict]):
    """Write structured items as DOCX paragraphs, preserving formatting (bold/italic).
    Handles both text items and image items."""
    for item in items:
        # Handle image items
        if item.get("type") == "image":
            try:
                img_url = item.get("url", "")
                alt_text = item.get("alt", "Image")
                if img_url:
                    # Download image
                    response = urlopen(img_url, timeout=5)
                    img_stream = BytesIO(response.read())
                    # Add centered image paragraph
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(img_stream, width=Inches(5))
                    log.info("  Embedded image: %s", alt_text)
            except Exception as e:
                log.debug("Could not embed image %s: %s", item.get("url", "?"), e)
            continue
        
        # Handle text items
        runs = item.get("runs", [])
        requested_style = item.get("style", "Normal")
        if not runs:
            continue
        
        # Get available style with fallback mappings
        style = _get_available_style(doc, requested_style)
        p = doc.add_paragraph(style=style)
        for i, run_data in enumerate(runs):
            text = run_data.get("text", "")
            bold = run_data.get("bold", False)
            italic = run_data.get("italic", False)
            if text:
                # Add space between runs if needed (unless at start or after space)
                if i > 0 and not text.startswith(" ") and not runs[i-1].get("text", "").endswith(" "):
                    p.add_run(" ")
                run = p.add_run(text)
                run.bold = bold
                run.italic = italic


def _copy_template_content(doc: Document, template_path: str) -> None:
    """Copy all content from template DOCX into the current document."""
    if not Path(template_path).exists():
        log.warning("Template file not found: %s", template_path)
        return
    
    try:
        template_doc = Document(template_path)
        
        # Copy all paragraphs from template
        for para in template_doc.paragraphs:
            new_para = doc.add_paragraph(style=para.style)
            # Copy paragraph properties
            new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
            new_para.paragraph_format.right_indent = para.paragraph_format.right_indent
            new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
            new_para.paragraph_format.space_before = para.paragraph_format.space_before
            new_para.paragraph_format.space_after = para.paragraph_format.space_after
            new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing
            new_para.alignment = para.alignment
            
            # Copy runs with formatting
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                if run.font.bold:
                    new_run.font.bold = True
                if run.font.italic:
                    new_run.font.italic = True
                if run.font.underline:
                    new_run.font.underline = True
                if run.font.size:
                    new_run.font.size = run.font.size
                if run.font.color.rgb:
                    new_run.font.color.rgb = run.font.color.rgb
        
        # Copy tables from template if present
        for table in template_doc.tables:
            tbl = table._tbl
            new_tbl = doc._element.add_tbl()
            for row in tbl.tr_lst:
                new_row = new_tbl.add_tr()
                for cell in row.tc_lst:
                    new_cell = new_row.add_tc()
                    new_cell._element.getparent().replace(cell, new_cell._element)
        
        log.info("Template content copied: %s", template_path)
    except Exception as e:
        log.error("Error copying template: %s", e)


def _replace_placeholders(doc: Document, title: str, base_url: str) -> None:
    """Replace placeholders in document: XXX1, HHH1, FFF1 with appropriate values.
    Searches in headers, main content, footers, and all runs."""
    
    # Replace XXX1 and HHH1 with title in header
    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            for run in para.runs:
                if "XXX1" in run.text or "HHH1" in run.text:
                    run.text = run.text.replace("XXX1", title).replace("HHH1", title)
    
    # Replace XXX1 and HHH1 in main document paragraphs (search all runs)
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text and ("XXX1" in run.text or "HHH1" in run.text):
                run.text = run.text.replace("XXX1", title).replace("HHH1", title)
    
    # Replace XXX1 and HHH1 in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text and ("XXX1" in run.text or "HHH1" in run.text):
                            run.text = run.text.replace("XXX1", title).replace("HHH1", title)
    
    # Replace FFF1 in footer (if template has it)
    for section in doc.sections:
        footer = section.footer
        footer_has_fff1 = False
        for para in footer.paragraphs:
            if "FFF1" in para.text:
                footer_has_fff1 = True
                break
        
        # Only recreate footer if FFF1 placeholder exists
        if footer_has_fff1:
            # Clear existing footer
            for para in footer.paragraphs:
                p = para._element
                p.getparent().remove(p)
            
            # Create new footer with proper formatting
            footer_para = footer.add_paragraph()
            footer_para.paragraph_format.tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
            footer_para.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
            
            # Left: URL as hyperlink in blue
            run_left = footer_para.add_run(base_url if base_url else "http://koyil.org")
            run_left.font.name = "Arial"
            run_left.font.size = Pt(10)
            run_left.font.color.rgb = RGBColor(0, 51, 204)
            run_left.font.underline = True
            
            footer_para.add_run("\t")
            
            # Center: page number
            run_page = footer_para.add_run()
            from docx.oxml import parse_xml
            fldChar1 = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))
            run_page._r.append(fldChar1)
            run = footer_para.add_run()
            instrText = parse_xml(r'<w:instrText {} xml:space="preserve">PAGE</w:instrText>'.format('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))
            run._r.append(instrText)
            run = footer_para.add_run()
            fldChar2 = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))
            run._r.append(fldChar2)
            
            run_num = footer_para.add_run("1")
            run_num.font.size = Pt(10)
            
            footer_para.add_run("\t")
            
            # Right: koyil.org in blue
            run_right = footer_para.add_run("koyil.org")
            run_right.font.name = "Arial"
            run_right.font.size = Pt(10)
            run_right.font.color.rgb = RGBColor(0, 51, 204)


def _load_template_path(language: str, cfg: dict) -> str:
    """Get the full path to the template file for the given language."""
    template_folder = cfg.get("template_folder", "")
    templates = cfg.get("templates", {})
    
    lang_lower = language.lower()
    template_name = templates.get(lang_lower)
    
    if not template_name:
        log.warning("Language '%s' not found in config. Available: %s", language, ", ".join(templates.keys()))
        return ""
    
    full_path = str(Path(template_folder) / template_name)
    return full_path


def build_docx(
    output_path,
    main_content_lines: list[str],
    chapters: list[dict],
    chapter_contents: list[dict],
    base_url: str = "",
    title: str = "",
    language: str = "english",
    cfg: dict = None,
) -> None:
    """
    Build and save the final .docx document in pure OOXML format.

    Document structure:
      Page 1  — Title + TOC (Word field — updates page numbers in MS Word)
      Page 2  — Main page scraped content (with inline images where they appear)
      Page 3+ — Chapter 1 Heading 2 + content (with inline images)
                 Chapter 2 Heading 2 + content (with inline images)
                 ...
    
    If language and cfg provided, loads template first and copies its content.
    Gracefully falls back to blank document if template is missing or invalid.
    """
    # Try to load template if language is provided
    doc = None
    if language and cfg:
        template_path = _load_template_path(language, cfg)
        if template_path and Path(template_path).exists():
            try:
                doc = Document(template_path)
                log.info("Starting with template for language: %s", language)
            except Exception as e:
                log.warning("Failed to load template %s: %s. Using blank document.", template_path, e)
                doc = None
        elif template_path:
            log.warning("Template file not found: %s. Using blank document.", template_path)
    
    # Create blank document if template loading failed or wasn't requested
    if doc is None:
        doc = Document()

    # If template was loaded, clear its header/footer to avoid duplication
    if language and cfg and doc is not None:
        _clear_all_headers_footers(doc)

    # Add header and footer FIRST (before adding content)
    # This ensures they apply correctly to the document
    _add_header(doc, title)
    _add_page_numbers(doc, base_url)

    # Pre-compute bookmark names for all chapters (needed in both TOC and headings)
    bookmark_names = [
        _make_bookmark_name(i, ch["title"])
        for i, ch in enumerate(chapters, 1)
    ]

    # -------------------------------------------------------------------
    # Page 1 — TOC
    # -------------------------------------------------------------------
    # Use Title style instead of Heading 1, so it doesn't appear in the auto-generated TOC
    toc_para = doc.add_paragraph("Table of Contents", style="Title")
    toc_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Embed a Word TOC field (auto-updates page numbers in MS Word with F9)
    _add_toc_field(doc)

    _add_page_break(doc)

    # -------------------------------------------------------------------
    # Page 2 — Main page content
    # -------------------------------------------------------------------
    doc.add_heading("Introduction", level=1)
    _write_lines(doc, main_content_lines)
    
    _add_page_break(doc)

    # -------------------------------------------------------------------
    # Pages 3+ — One section per chapter, each heading carries a bookmark
    # -------------------------------------------------------------------
    for i, ch_data in enumerate(chapter_contents):
        title = ch_data["title"]
        lines = ch_data["lines"]
        bm_name = bookmark_names[i]
        bm_id   = i + 1  # IDs start at 1

        # Use Heading 1 for chapters (same level as Introduction) so TOC alignment is consistent
        # Add sequence number to chapter title for TOC numbering
        numbered_title = f"{i + 1}. {title}"
        hdg = doc.add_heading(numbered_title, level=1)
        # Make heading blue
        for run in hdg.runs:
            run.font.color.rgb = RGBColor(0, 51, 204)  # Blue color
        _add_bookmark(hdg, bm_id, bm_name)

        if lines:
            _write_lines(doc, lines)
        else:
            doc.add_paragraph("[No content extracted for this chapter.]", style="Normal")

        _add_page_break(doc)

    # -------------------------------------------------------------------
    # Replace placeholders if template was used
    # -------------------------------------------------------------------
    if language and cfg:
        _replace_placeholders(doc, title, base_url)

    # -------------------------------------------------------------------
    # Save (delete existing file first)
    # -------------------------------------------------------------------
    # Delete existing file if it exists (can't overwrite while locked)
    path_obj = Path(output_path)
    if path_obj.exists():
        try:
            path_obj.unlink()
            log.info("Deleted existing file: %s", output_path)
        except OSError as exc:
            log.warning("Could not delete existing file: %s", exc)

    # Save document
    doc.save(str(output_path))
    log.info("Document saved: %s", output_path)
    log.info("Updating Word fields for TOC page numbers...")
    _update_word_fields_async(str(output_path))


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------
def main():
    args = parse_args()
    cfg = load_config()

    # Validate inputs before doing anything expensive
    url = validate_url(args.url)
    output_path = validate_output_path(args.output)
    title = args.title
    language = args.language

    log.info("Input URL   : %s", url)
    log.info("Output file : %s", output_path)
    log.info("Document title: %s", title)
    log.info("Template language: %s", language)

    # Clean up old DOCX files in the working directory
    from pathlib import Path
    output_dir = Path(output_path).parent
    for old_file in output_dir.glob("*.docx"):
        if old_file.name != Path(output_path).name:
            try:
                old_file.unlink()
                log.info("Deleted old file: %s", old_file.name)
            except OSError as e:
                log.debug("Could not delete %s: %s", old_file.name, e)

    with sync_playwright() as pw:
        browser, page = open_page(pw, url, cfg)

        # Task 3 — extract main page content
        main_html = get_page_html(page)
        content_items = scrape_page(page, cfg, url)

        if content_items:
            log.info("Main page: extracted %d items.", len(content_items))
        else:
            log.warning("No content extracted from main page.")

        # Task 4 — detect chapters (TOC)
        toc_selector = cfg.get("toc_selector", "li > a[href]")
        max_chapters = cfg.get("max_chapters", 100)
        chapters = detect_chapters(main_html, url, toc_selector, max_chapters)

        if chapters:
            log.info("--- TOC preview (first 10 chapters) ---")
            for i, ch in enumerate(chapters[:10], 1):
                log.info("  %02d. %s", i, ch["title"])
            if len(chapters) > 10:
                log.info("  ... and %d more chapters", len(chapters) - 10)
            log.info("--- (total %d chapters) ---", len(chapters))
        else:
            log.warning("No chapter links detected on main page.")

        # Task 4.4 — scrape each chapter page
        chapter_contents = []   # list of {"title": str, "lines": list[dict]}
        for idx, ch in enumerate(chapters, 1):
            log.info("Scraping chapter %d/%d: %s", idx, len(chapters), ch["title"])
            try:
                page.goto(ch["url"], wait_until="networkidle")
            except PlaywrightTimeout:
                log.warning("Timeout on chapter page: %s — skipping.", ch["url"])
                chapter_contents.append({"title": ch["title"], "lines": []})
                continue

            ch_items = scrape_page(page, cfg, ch["url"])
            log.info("  -> %d items extracted.", len(ch_items))
            chapter_contents.append({"title": ch["title"], "lines": ch_items})

        # Task 5 — build and save the DOCX
        log.info("Building DOCX document ...")
        build_docx(output_path, content_items, chapters, chapter_contents, url, title, language, cfg)

        browser.close()

    log.info("Done.")


if __name__ == "__main__":
    main()

